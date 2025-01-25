# Custom Invitations as Word Documents
# Takes guest names from a txt file, generates a Word document with custom invitations.

import docx
import logging

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')


def get_list_from_txt_file(filename: str) -> list:
    try:
        with open(filename) as txt_file:
            txt_file_contents = txt_file.read()
            return txt_file_contents.split('\n')
    except Exception as e:
        logger.error(f'Could not open the file: {filename}, error: {e}')


def generate_custom_invitations(docx_template_filename: str, guests: list):
    doc = docx.Document(docx_template_filename)
    for guest in guests:
        par1 = doc.add_paragraph('It would be a pleasure to have the company of')
        par1.style = 'Beautiful Style'
        par2 = doc.add_paragraph(guest)
        par2.style = 'Guest Name'
        par3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
        par3.style = 'Beautiful Style'
        par4 = doc.add_paragraph('April 1st')
        par4.style = 'Invitation Date'
        par5 = doc.add_paragraph("at 7 o'clock")
        par5.style = 'Beautiful Style'
        par5.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    new_filename = 'invitations.docx'
    doc.save(new_filename)
    logger.info(f'The file {new_filename} has been created!')


if __name__ == '__main__':
    guests_list = get_list_from_txt_file('guests.txt')
    generate_custom_invitations('template.docx', guests_list)
